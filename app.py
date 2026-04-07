import streamlit as st
import pandas as pd
from pypdf import PdfReader
from docx import Document
import io
import litellm
import os
import google.generativeai as genai

# --- ASSET LOADING ---
def load_css(file_path):
    with open(file_path, encoding='utf-8') as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

def load_html(file_path):
    with open(file_path, encoding='utf-8') as f:
        st.markdown(f.read(), unsafe_allow_html=True)

# Import Prompt from assets
try:
    from assets.prompts import SYSTEM_PROMPT
except ImportError:
    st.error("Error: assets/prompts.py not found.")
    SYSTEM_PROMPT = "You are a helpful assistant."

# Page Configuration
st.set_page_config(page_title="AI Document Intelligence", page_icon="📄", layout="wide")

# Load external CSS
if os.path.exists("assets/style.css"):
    load_css("assets/style.css")

# --- UTILITY FUNCTIONS ---

def extract_text_from_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            text += " | ".join([cell.text for cell in row.cells]) + "\n"
    return text

def extract_text_from_excel(file):
    df = pd.read_excel(file)
    return df.to_string(index=False)

def get_context(uploaded_files):
    context = ""
    for file in uploaded_files:
        if file.name.endswith('.pdf'):
            context += f"--- Filename: {file.name} ---\n" + extract_text_from_pdf(file) + "\n\n"
        elif file.name.endswith('.docx'):
            context += f"--- Filename: {file.name} ---\n" + extract_text_from_docx(file) + "\n\n"
        elif file.name.endswith('.xlsx') or file.name.endswith('.xls'):
            context += f"--- Filename: {file.name} ---\n" + extract_text_from_excel(file) + "\n\n"
    return context

# --- SIDEBAR CONFIGURATION ---
with st.sidebar:
    # Modern Enterprise Logo
    st.markdown("""
    <div style='display: flex; align-items: center; gap: 10px; margin-bottom: 25px;'>
        <div style='width: 32px; height: 32px; background: #0f172a; border-radius: 8px; display: flex; justify-content: center; align-items: center;'>
            <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="white" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="18" height="18" rx="2" ry="2"></rect><line x1="9" y1="3" x2="9" y2="21"></line></svg>
        </div>
        <span style='font-weight: 800; font-size: 1.1rem; letter-spacing: 0.05em; color: #0f172a;'>NEXUS UI</span>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<p style='font-size: 0.75rem; color: #475569; font-weight: 700; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 5px;'>System Configuration</p>", unsafe_allow_html=True)

    # Provider Selection
    provider = st.selectbox(
        "Select Provider", 
        ["Google Gemini", "OpenAI", "Anthropic", "Perplexity"],
        index=0
    )
    
    api_key_input = st.text_input(f"Enter {provider} API Key", type="password")
    
    # Suggested Models based on Provider (Updated for Stability & Capability)
    model_options = {
        "Google Gemini": ["gemini/gemini-flash-latest", "gemini/gemini-2.5-flash", "gemini/gemini-2.5-pro", "gemini/gemini-pro-latest"],
        "OpenAI": ["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"],
        "Anthropic": ["anthropic/claude-3-5-sonnet-20240620", "anthropic/claude-3-opus-20240229", "anthropic/claude-3-haiku-20240307"],
        "Perplexity": ["perplexity/llama-3-sonar-large-32k-online", "perplexity/llama-3-sonar-small-32k-online"]
    }
    
    model_choice = st.selectbox("Select Model", model_options.get(provider, ["gpt-3.5-turbo"]))
    
    if not api_key_input:
        st.warning(f"Please enter your {provider} API Key.")
    else:
        st.success("API Key configured!")
            
    st.markdown("---")
    st.markdown("### 📄 Knowledge Base")
    uploaded_files = st.file_uploader(
        "Upload PDF, DOCX, or XLSX", 
        type=["pdf", "docx", "xlsx", "xls"], 
        accept_multiple_files=True,
        key="doc_uploader"
    )

    if uploaded_files:
        st.success(f"Loaded {len(uploaded_files)} source files")

    st.markdown("---")
    if st.button("🗑️ Clear Conversation", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

    st.markdown("---")
    st.markdown(f"📡 **Model:** `{model_choice}`")

# --- CHAT INTERFACE ---

# Load Header HTML
if os.path.exists("assets/header.html"):
    load_html("assets/header.html")
else:
    st.title("AI Document Intelligence")

if "messages" not in st.session_state:
    st.session_state.messages = []

# Display chat history
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Chat Input
if question := st.chat_input("Input your query..."):
    if not api_key_input:
        st.error(f"Please enter your {provider} API Key in the sidebar.")
    elif not uploaded_files:
        st.warning("Please upload some documents first.")
    else:
        # User message
        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        # AI response
        with st.chat_message("assistant"):
            try:
                # 1. Gather Context
                with st.status("🔍 Scanning documents...", expanded=False) as status:
                    context = get_context(uploaded_files)
                    status.update(label="✅ Context extracted!", state="complete")
                
                # 2. Build Prompt
                full_prompt = f"{SYSTEM_PROMPT}\n\nContext:\n{context}\n\nUser Question:\n{question}"
                
                # 3. Get Response
                with st.spinner(f"Thinking (via {provider})..."):
                    if provider == "Google Gemini":
                        genai.configure(api_key=api_key_input)
                        native_model = model_choice.replace("gemini/", "")
                        try:
                            model = genai.GenerativeModel(native_model)
                            response = model.generate_content(full_prompt)
                            answer = response.text
                        except Exception as gemini_err:
                            st.error(f"❌ Google API declined the request for '{native_model}'.")
                            try:
                                allowed = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                                st.info("💡 Diagnostic: According to Google, your specific API Key only has access to these active models:\n" + "\n".join([f"- `{m}`" for m in allowed]))
                            except Exception:
                                pass
                            st.error(f"Raw Error: {str(gemini_err)}")
                            st.stop()
                    else:
                        completion_kwargs = {
                            "model": model_choice,
                            "messages": [{"role": "user", "content": full_prompt}],
                            "api_key": api_key_input,
                            "temperature": 0.5,
                            "max_tokens": None
                        }
                        response = litellm.completion(**completion_kwargs)
                        answer = response.choices[0].message.content
                
                st.markdown(answer)
                st.session_state.messages.append({"role": "assistant", "content": answer})
                    
            except Exception as e:
                st.error(f"Error: {str(e)}")
